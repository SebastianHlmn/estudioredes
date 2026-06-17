"""
reports/report_generator.py
Exportación inicial a Word.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document

from config import EXPORTS_DIR


def add_dataframe_table(document: Document, df: pd.DataFrame, title: str, max_rows: int = 20) -> None:
    document.add_heading(title, level=2)
    if df.empty:
        document.add_paragraph("Sin datos disponibles.")
        return
    view = df.head(max_rows).copy()
    table = document.add_table(rows=1, cols=len(view.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(view.columns):
        hdr_cells[i].text = str(col)
    for _, row in view.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(view.columns):
            value = row[col]
            cells[i].text = "" if pd.isna(value) else str(value)


def generate_basic_report(
    project_name: str,
    posts_df: pd.DataFrame,
    concept_summary_df: pd.DataFrame,
    output_dir: Path | None = None,
) -> Path:
    output_dir = output_dir or EXPORTS_DIR
    output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    output_path = output_dir / f"reporte_estudioredes_{timestamp}.docx"

    doc = Document()
    doc.add_heading(f"Reporte inicial - {project_name}", level=1)
    doc.add_paragraph(
        "Este reporte resume el estado inicial del relevamiento, la clasificación discursiva y la trayectoria de conceptos detectados."
    )

    doc.add_heading("Resumen", level=2)
    doc.add_paragraph(f"Total de publicaciones normalizadas: {len(posts_df)}")

    if not posts_df.empty and "platform" in posts_df.columns:
        platform_df = posts_df.groupby("platform").size().reset_index(name="publicaciones")
        add_dataframe_table(doc, platform_df, "Publicaciones por plataforma")

    if not posts_df.empty and "frame" in posts_df.columns:
        frame_df = posts_df.groupby("frame", dropna=False).size().reset_index(name="publicaciones")
        frame_df = frame_df.sort_values("publicaciones", ascending=False)
        add_dataframe_table(doc, frame_df, "Marcos discursivos detectados")

    add_dataframe_table(doc, concept_summary_df, "Trayectoria y cristalización de conceptos")

    sample_cols = [c for c in ["created_at", "platform", "text", "frame", "discursive_strategy", "concept"] if c in posts_df.columns]
    if sample_cols:
        add_dataframe_table(doc, posts_df[sample_cols], "Muestra de publicaciones", max_rows=10)

    doc.add_heading("Nota metodológica", level=2)
    doc.add_paragraph(
        "La versión inicial utiliza un clasificador interpretable basado en reglas y un codebook auditable. "
        "Las siguientes iteraciones incorporarán modelos de lenguaje, validación humana sistemática, embeddings semánticos y métricas de red más avanzadas."
    )

    doc.save(output_path)
    return output_path

# fin reports/report_generator.py
